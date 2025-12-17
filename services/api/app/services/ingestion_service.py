"""
Loan document ingestion - handles PDF/DOCX parsing and data extraction
TODO: Replace regex patterns with proper NLP/ML extraction for production
"""
import uuid
from datetime import datetime
from typing import Dict, Any, List
from pathlib import Path
import PyPDF2
from docx import Document
import re

from app.models import Loan, LoanDocument, Covenant, ESGClause


class IngestionService:
    """Handles document parsing and extraction - basic implementation for hackathon"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']
        # TODO: Add support for scanned PDFs (OCR) if needed
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF - PyPDF2 works for most cases but struggles with complex layouts"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                # Some PDFs have empty pages, skip them
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text.strip()
        except Exception as e:
            # Re-raise with more context
            raise ValueError(f"Failed to read PDF {file_path}: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text based on file type"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def extract_covenants(self, text: str) -> List[Covenant]:
        """Extract covenants using regex - works for demo but needs ML for real docs"""
        covenants = []
        
        # These patterns catch most common covenant formats
        # Real-world: would use spaCy/NER or fine-tuned BERT for this
        covenant_patterns = [
            r"(?:debt[\s-]?to[\s-]?equity|D/E|debt[\s-]?equity)[\s:]+(?:ratio|must|shall)?[\s:]*([<>≤≥=]+)?[\s]*([\d.]+)",
            r"(?:current[\s-]?ratio)[\s:]+(?:must|shall)?[\s:]*([<>≤≥=]+)?[\s]*([\d.]+)",
            r"(?:interest[\s-]?coverage)[\s:]+(?:ratio|must|shall)?[\s:]*([<>≤≥=]+)?[\s]*([\d.]+)",
            r"(?:minimum[\s-]?equity)[\s:]+(?:must|shall)?[\s:]*([<>≤≥=]+)?[\s]*([\d.]+)",
        ]
        
        covenant_names = ["Debt-to-Equity", "Current Ratio", "Interest Coverage", "Minimum Equity"]
        for i, pattern in enumerate(covenant_patterns):
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                # Handle missing operator gracefully
                operator = match.group(1) if match.group(1) else ">="
                threshold_str = match.group(2) if match.group(2) else None
                
                if not threshold_str:
                    continue  # Skip if we can't extract threshold
                
                try:
                    threshold = float(threshold_str)
                except (ValueError, TypeError):
                    continue  # Skip malformed numbers
                
                covenant = Covenant(
                    id=str(uuid.uuid4()),
                    name=covenant_names[i] if i < len(covenant_names) else f"Covenant_{i+1}",
                    type="financial",
                    threshold=threshold,
                    operator=operator,
                    frequency="quarterly",  # Default - could parse from doc
                    next_check_date=datetime.now(),
                    description="Extracted from document"
                )
                covenants.append(covenant)
        
        # Fallback for demo - ensures we always have something to show
        if not covenants:
            covenants.append(Covenant(
                id=str(uuid.uuid4()),
                name="Default Financial Covenant",
                type="financial",
                threshold=2.0,
                operator=">=",
                frequency="quarterly",
                next_check_date=datetime.now(),
                description="Default covenant for demonstration"
            ))
        
        return covenants
    
    def extract_esg_clauses(self, text: str) -> List[ESGClause]:
        """Extract ESG clauses from document text"""
        esg_clauses = []
        
        # Pattern matching for ESG-related content
        esg_keywords = {
            "environmental": ["carbon", "emission", "sustainability", "environment", "green"],
            "social": ["diversity", "labor", "safety", "community", "social"],
            "governance": ["compliance", "ethics", "board", "transparency", "governance"]
        }
        
        for category, keywords in esg_keywords.items():
            found = any(keyword.lower() in text.lower() for keyword in keywords)
            if found:
                clause = ESGClause(
                    id=str(uuid.uuid4()),
                    category=category,
                    requirement=f"{category.title()} compliance required",
                    reporting_frequency="annually",
                    next_report_date=datetime.now(),
                    description=f"ESG {category} requirement extracted from document"
                )
                esg_clauses.append(clause)
        
        # Default ESG clause if none found
        if not esg_clauses:
            esg_clauses.append(ESGClause(
                id=str(uuid.uuid4()),
                category="governance",
                requirement="General governance compliance",
                reporting_frequency="annually",
                next_report_date=datetime.now(),
                description="Default ESG clause for demonstration"
            ))
        
        return esg_clauses
    
    def extract_loan_metadata(self, text: str) -> Dict[str, Any]:
        """Extract basic loan metadata - amount, rate, etc."""
        metadata = {}
        
        # Try multiple patterns for loan amount - documents vary a lot
        amount_patterns = [
            r"(?:loan[\s-]?amount|principal|amount)[\s:]+(?:is|of)?[\s:]*\$?([\d,]+\.?\d*)",
            r"\$([\d,]+\.?\d*)[\s]*(?:loan|principal)",
            r"(?:principal|loan)[\s:]+(?:of|is)?[\s:]*\$?([\d,]+\.?\d*)"
        ]
        for pattern in amount_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                amount_str = match.group(1).replace(',', '').replace('$', '')
                try:
                    metadata['loan_amount'] = float(amount_str)
                    break
                except (ValueError, AttributeError):
                    continue
        
        # Interest rate extraction
        rate_pattern = r"(?:interest[\s-]?rate|rate)[\s:]+(?:is|of)?[\s:]*([\d.]+)[\s]*(?:%|percent|basis[\s-]?points)?"
        match = re.search(rate_pattern, text, re.IGNORECASE)
        if match:
            try:
                rate = float(match.group(1))
                # Sanity check - if it looks like basis points, convert
                if rate > 100:
                    rate = rate / 100
                metadata['interest_rate'] = rate
            except (ValueError, AttributeError):
                pass
        
        # Sensible defaults for demo
        metadata.setdefault('loan_amount', 1000000.0)
        metadata.setdefault('interest_rate', 5.5)
        
        return metadata
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Process a loan document and extract structured data
        
        Returns:
            Dictionary with extracted loan data, covenants, and ESG clauses
        """
        # Extract text
        text = self.extract_text(file_path)
        
        # Extract components
        covenants = self.extract_covenants(text)
        esg_clauses = self.extract_esg_clauses(text)
        metadata = self.extract_loan_metadata(text)
        
        # Create loan document record
        loan_doc = LoanDocument(
            id=str(uuid.uuid4()),
            loan_id="",  # Will be set when loan is created
            filename=filename,
            file_type=Path(file_path).suffix[1:].lower(),
            upload_date=datetime.now(),
            extracted_data={
                "text_length": len(text),
                "covenant_count": len(covenants),
                "esg_clause_count": len(esg_clauses)
            }
        )
        
        return {
            "document": loan_doc,
            "covenants": covenants,
            "esg_clauses": esg_clauses,
            "metadata": metadata,
            "extracted_text": text[:1000]  # First 1000 chars for reference
        }

